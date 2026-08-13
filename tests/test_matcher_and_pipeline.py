"""Evaluation matcher behaviour and an end-to-end pipeline smoke test."""
import os

import docx

from pii_redactor.config import load_settings
from pii_redactor.document.loader import DocxDocument
from pii_redactor.evaluation.matcher import match_unit
from pii_redactor.models import GoldSpan, Span
from pii_redactor.redactor import redact

CONFIG = os.path.join(os.path.dirname(__file__), "..", "config", "default.yaml")


def test_matcher_strict_vs_overlap():
    gold = [GoldSpan("u", 0, 10, "NAME")]
    pred = [Span(2, 8, "x", "NAME", "d")]          # inside gold, different offsets
    tp_s, _, _ = match_unit(pred, gold, "strict")
    tp_o, _, _ = match_unit(pred, gold, "overlap")
    assert tp_s == 0 and tp_o == 1


def test_matcher_one_to_one():
    gold = [GoldSpan("u", 0, 5, "PHONE"), GoldSpan("u", 6, 11, "PHONE")]
    pred = [Span(0, 5, "a", "PHONE", "d")]
    tp, mp, mg = match_unit(pred, gold, "overlap")
    assert tp == 1 and len(mp) == 1


def test_pipeline_end_to_end(tmp_path):
    src = tmp_path / "in.docx"
    d = docx.Document()
    d.add_paragraph("Contact Person: Rashi Patil")
    d.add_paragraph("Email: rashi.patil@gmail.com Telephone: +91 98765 43210")
    d.save(src)

    out_path = tmp_path / "out.docx"
    settings = load_settings(CONFIG)
    result = redact(str(src), str(out_path), settings)

    red = "\n".join(u.text for u in DocxDocument(str(out_path)).units())
    assert "rashi.patil@gmail.com" not in red      # email redacted
    assert "98765 43210" not in red                # phone redacted
    assert "@example.com" in red                   # replaced with a fake
    # audit contains no raw PII
    assert all("gmail" not in e["replacement"] for e in result.auditor.events)
