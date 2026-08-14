"""The web service must reuse the existing pipeline and return usable bytes."""
import io

import docx

from pii_redactor.document.loader import DocxDocument
from pii_redactor.webservice import redact_bytes


def _sample_docx_bytes() -> bytes:
    d = docx.Document()
    d.add_paragraph("Contact Person: Rashi Patil")
    d.add_paragraph("Email: rashi.patil@gmail.com Telephone: +91 98765 43210")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_redact_bytes_returns_summary_and_files():
    out = redact_bytes(_sample_docx_bytes(), ["EMAIL", "PHONE", "NAME"])
    assert out["summary"]["total_redactions"] >= 2
    assert out["summary"]["by_category"].get("EMAIL", 0) >= 1
    assert out["redacted_bytes"][:2] == b"PK"          # valid .docx (zip) header
    assert out["audit_bytes"]                            # audit log produced


def test_redact_bytes_respects_category_selection(tmp_path):
    # only PHONE selected -> email must survive
    out = redact_bytes(_sample_docx_bytes(), ["PHONE"])
    p = tmp_path / "r.docx"
    p.write_bytes(out["redacted_bytes"])
    text = "\n".join(u.text for u in DocxDocument(str(p)).units())
    assert "rashi.patil@gmail.com" in text              # email not redacted
    assert "98765 43210" not in text                    # phone redacted
    assert "EMAIL" not in out["summary"]["by_category"]


def test_redact_bytes_no_raw_pii_in_audit():
    out = redact_bytes(_sample_docx_bytes(), ["EMAIL", "PHONE", "NAME"])
    assert b"gmail" not in out["audit_bytes"]
    assert b"Rashi" not in out["audit_bytes"]
