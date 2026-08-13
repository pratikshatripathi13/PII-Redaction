"""Determinism of pseudonymization, no-raw-PII audit, and run-level rewriting."""
import docx

from pii_redactor.document.loader import DocxDocument
from pii_redactor.document.rewriter import apply_replacements
from pii_redactor.models import Span
from pii_redactor.pseudonym.mapping_store import MappingStore


def test_same_entity_same_replacement():
    m = MappingStore(salt="test")
    a = m.replacement_for("NAME", "Kushal Subbayya Hegde")
    b = m.replacement_for("NAME", "Kushal Subbayya Hegde")
    assert a == b


def test_normalization_equivalence():
    m = MappingStore(salt="test")
    # different spacing / casing of the same phone/email map to the same fake
    assert m.replacement_for("PHONE", "+91 22 40094400") == m.replacement_for("PHONE", "+912240094400")
    assert m.replacement_for("EMAIL", "A@B.COM") == m.replacement_for("EMAIL", "a@b.com")


def test_distinct_entities_distinct_replacements():
    m = MappingStore(salt="test")
    assert m.replacement_for("NAME", "Alpha Beta") != m.replacement_for("NAME", "Gamma Delta")


def test_replacement_reproducible_across_stores():
    # cross-run determinism: a fresh store with the same salt yields the same value
    a = MappingStore(salt="s").replacement_for("EMAIL", "x@y.com")
    b = MappingStore(salt="s").replacement_for("EMAIL", "x@y.com")
    assert a == b


def test_audit_id_is_hashed_not_raw():
    m = MappingStore(salt="s")
    hid = m.hashed_id("EMAIL", "secret.person@corp.com")
    assert "secret" not in hid and "@" not in hid and len(hid) == 16


def test_email_format_preserved():
    m = MappingStore(salt="s")
    assert m.replacement_for("EMAIL", "a.b@c.com").endswith("@example.com")


def test_rewrite_multi_run_span(tmp_path):
    # Build a paragraph whose email is spread across runs incl. empty runs,
    # then confirm the replacement is written exactly once (regression test).
    d = docx.Document()
    p = d.add_paragraph()
    for t in ["Email: ", "", "cs.connect@", "", "ksh.com", " end"]:
        p.add_run(t)
    fp = tmp_path / "t.docx"
    d.save(fp)

    doc = DocxDocument(str(fp))
    unit = next(u for u in doc.units() if "cs.connect@ksh.com" in u.text)
    text = unit.text
    s = text.index("cs.connect@ksh.com")
    span = Span(s, s + len("cs.connect@ksh.com"), "cs.connect@ksh.com", "EMAIL", "t")
    apply_replacements(unit, [(span, "john.doe@example.com")])
    new = "".join(r.text for r in unit.runs)
    assert new == "Email: john.doe@example.com end"
    assert new.count("example.com") == 1  # no duplication from empty runs


def test_rewrite_preserves_run_formatting(tmp_path):
    d = docx.Document()
    p = d.add_paragraph()
    r1 = p.add_run("Bold name ")
    r1.bold = True
    r2 = p.add_run("Kushal Hegde")
    r2.italic = True
    fp = tmp_path / "f.docx"
    d.save(fp)

    doc = DocxDocument(str(fp))
    unit = next(u for u in doc.units() if "Kushal Hegde" in u.text)
    text = unit.text
    s = text.index("Kushal Hegde")
    span = Span(s, s + len("Kushal Hegde"), "Kushal Hegde", "NAME", "t")
    apply_replacements(unit, [(span, "Fake Person")])
    runs = unit.runs
    assert runs[0].bold is True                     # untouched run keeps bold
    assert runs[1].italic is True                    # replaced run keeps italic
    assert "".join(r.text for r in runs) == "Bold name Fake Person"

