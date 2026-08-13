"""DOCX document model.

A `TextUnit` is one paragraph (top-level, inside a table cell, or inside a
header/footer). Each unit exposes:
  * text        - the concatenated visible text of its runs
  * runs        - the underlying python-docx Run objects
  * run_offsets - (start, end) char range of each run within `text`

Detection and the run-level rewriter both operate on this single serialization,
so character offsets always line up between them and between the evaluator.
"""
from __future__ import annotations

from dataclasses import dataclass
from typing import Iterator

import docx
from docx.document import Document as _Doc
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class TextUnit:
    unit_id: str
    paragraph: Paragraph
    kind: str  # "body" | "cell" | "header" | "footer"

    @property
    def runs(self):
        return self.paragraph.runs

    @property
    def text(self) -> str:
        return "".join(r.text for r in self.paragraph.runs)

    @property
    def run_offsets(self) -> list[tuple[int, int]]:
        offsets, pos = [], 0
        for r in self.paragraph.runs:
            n = len(r.text)
            offsets.append((pos, pos + n))
            pos += n
        return offsets


class DocxDocument:
    def __init__(self, path: str):
        self.path = path
        self.doc: _Doc = docx.Document(path)

    # -- iteration ---------------------------------------------------------
    def units(self) -> Iterator[TextUnit]:
        idx = 0
        # Merged table cells make python-docx hand back the SAME paragraph element
        # more than once. De-duplicate by the element's stable XPath (NOT id(), whose
        # value is recycled across transient lxml proxies and mis-fires). This ensures
        # every paragraph is visited and redacted exactly once.
        seen: set[str] = set()

        def _new(p) -> bool:
            el = p._p
            key = el.getroottree().getpath(el)
            if key in seen:
                return False
            seen.add(key)
            return True

        # 1) body paragraphs and tables, in document order
        for block in self._iter_body_blocks(self.doc):
            if isinstance(block, Paragraph):
                if _new(block):
                    yield TextUnit(f"body-{idx}", block, "body"); idx += 1
            elif isinstance(block, Table):
                for p in self._iter_table_paragraphs(block):
                    if _new(p):
                        yield TextUnit(f"cell-{idx}", p, "cell"); idx += 1
        # 2) headers & footers (empty in this document, handled defensively)
        for si, section in enumerate(self.doc.sections):
            for p in section.header.paragraphs:
                if _new(p):
                    yield TextUnit(f"hdr-{si}-{idx}", p, "header"); idx += 1
            for p in section.footer.paragraphs:
                if _new(p):
                    yield TextUnit(f"ftr-{si}-{idx}", p, "footer"); idx += 1

    def _iter_body_blocks(self, parent):
        """Yield Paragraph and Table objects in true document order."""
        body = parent.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, parent)
            elif child.tag == qn("w:tbl"):
                yield Table(child, parent)

    def _iter_table_paragraphs(self, table: Table) -> Iterator[Paragraph]:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nested in cell.tables:  # nested tables
                    yield from self._iter_table_paragraphs(nested)

    def save(self, out_path: str):
        self.doc.save(out_path)
